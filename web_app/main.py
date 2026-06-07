from __future__ import annotations

import json
import os
import re
import shutil
import sys
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any

import httpx
from fastapi import FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from dotenv import load_dotenv
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

ROOT_DIR = Path(__file__).resolve().parents[1]
if str(ROOT_DIR) not in sys.path:
    sys.path.insert(0, str(ROOT_DIR))

from app import (  # noqa: E402
    build_gap_memory_content,
    build_gap_lines,
    build_rag_query,
    build_strength_lines,
    build_suggestion_lines,
    get_score_comment,
    join_skill_names,
    truncate_text,
)
from database import save_analysis_results  # noqa: E402
from tools import (  # noqa: E402
    AgentTraceStep,
    JobMemoryTool,
    JobRAGTool,
    build_agent_trace_step,
    compute_match_score,
    extract_candidate_skills,
    extract_job_requirements,
    get_jd_text,
    load_text_file,
)


WEB_DIR = Path(__file__).resolve().parent
APP_NAME = "Job Fit Agent Web"
APP_SLUG = "job-fit-agent"
DEFAULT_HOST = "127.0.0.1"
DEFAULT_PORT = 8000
RESULTS_DIR = WEB_DIR / "处理结果"
AI_CONFIG_DIR = WEB_DIR / "ai_config"
RESULTS_DIR.mkdir(exist_ok=True)
AI_CONFIG_DIR.mkdir(exist_ok=True)

app = FastAPI(title=APP_NAME)
app.mount("/static", StaticFiles(directory=WEB_DIR / "static"), name="static")
templates = Jinja2Templates(directory=WEB_DIR / "templates")


@app.middleware("http")
async def add_no_cache_headers(request: Request, call_next: Any) -> Any:
    response = await call_next(request)
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    return response


AI_SYSTEM_PROMPT = """
你是一个面向 AI 应用开发、后端 API、RAG 和 Agent 工程岗位的简历筛选复核助手。
你只基于本地规则已经算好的结构化结果做专业化表达，不要编造简历里没有的经历。
必须保留原始 filename、score、level，不重新排序、不重新打分。
只输出 JSON 对象，不要输出 Markdown、标题、解释、代码块或额外文本。
输出格式：
{
  "summary": "一句话说明 AI 已完成结构化复核",
  "reviews": [
    {
      "filename": "必须与输入完全一致",
      "conclusion": "一句话结论",
      "strengths": ["优势 1", "优势 2"],
      "gaps": ["短板 1", "短板 2"],
      "suggestions": ["建议 1", "建议 2", "建议 3"]
    }
  ]
}
每个数组保留 2 到 4 条，每条尽量控制在 18 到 60 个中文字符。
表达要像技术面试筛选意见，突出 API 接入、服务化、RAG、工具调用、部署和工程协作。
""".strip()
AI_REPORT_SYSTEM_PROMPT = """
你是一个面向招聘负责人和技术面试官的 AI 应用开发岗位筛选报告助手。
你会基于系统已经算好的简历 JD 匹配结果，写一份可直接交付的中文汇总报告。
不要编造简历里没有的经历，不要修改候选人的分数和等级。
报告要有专业 API/LLM 工程评审感，重点关注 FastAPI 服务化、大模型 API 接入、RAG、向量检索、Function Calling、Docker 部署和工程协作。
请输出 Markdown 正文，不要输出代码块。
""".strip()
AI_TIMEOUT_SECONDS = 25
AI_MAX_TOKENS = 1800
AI_REPORT_MAX_TOKENS = 3200
AI_RESUME_TEXT_LIMIT = 1200
AI_JD_TEXT_LIMIT = 1500
DEFAULT_WEB_USER_ID = "web_user"


def safe_filename(filename: str | None) -> str:
    name = Path(filename or "uploaded_file").name
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", name).strip(" .")
    return cleaned or "uploaded_file"


def save_upload(upload: UploadFile, target_dir: Path) -> Path:
    target = target_dir / safe_filename(upload.filename)
    with target.open("wb") as output:
        shutil.copyfileobj(upload.file, output)
    return target


def get_optional_env(*names: str) -> str | None:
    for name in names:
        value = os.environ.get(name)
        if value and value.strip():
            return value.strip()
    return None


def load_ai_config() -> dict[str, str] | None:
    load_dotenv(ROOT_DIR / ".env")
    load_dotenv(WEB_DIR / ".env", override=True, encoding="utf-8-sig")

    provider = get_optional_env("WEB_AI_PROVIDER", "AI_PROVIDER") or "openai"
    api_key = get_optional_env(
        "WEB_AI_API_KEY",
        "OPENAI_API_KEY",
        "DEEPSEEK_API_KEY",
        "ANTHROPIC_AUTH_TOKEN",
    )
    base_url = get_optional_env(
        "WEB_AI_BASE_URL",
        "OPENAI_BASE_URL",
        "DEEPSEEK_BASE_URL",
        "ANTHROPIC_BASE_URL",
    )
    model = get_optional_env(
        "WEB_AI_MODEL",
        "MODEL_NAME",
        "DEEPSEEK_MODEL",
        "ANTHROPIC_MODEL",
    )
    if not api_key or not base_url or not model:
        return None

    return {
        "provider": provider.strip().lower(),
        "api_key": api_key,
        "base_url": base_url.rstrip("/"),
        "model": model,
    }


def build_batch_ai_prompt(results: list[dict[str, Any]], jd_text: str) -> str:
    safe_jd_text = jd_text[:AI_JD_TEXT_LIMIT]
    compact_results = []
    for item in results:
        if not item.get("ok"):
            compact_results.append(
                {
                    "filename": item.get("filename"),
                    "error": item.get("error"),
                }
            )
            continue

        compact_results.append(
            {
                "filename": item["filename"],
                "score": item["score"],
                "level": item["level"],
                "conclusion": item["conclusion"],
                "matched_skills": item["matched_skills"],
                "missing_skills": item["missing_skills"],
                "strengths": item["strengths"],
                "gaps": item["gaps"],
                "suggestions": item["suggestions"],
            }
        )

    return f"""
岗位 JD：
{safe_jd_text}

本地规则批量分析结果：
{json.dumps(compact_results, ensure_ascii=False, indent=2)}

请把每个候选人的卡片内容做一次专业化复核改写。
硬性要求：
- 不要写候选人整体排序、第一档/第二档、最适合进入下一轮等批量报告式内容。
- 不要改变 score、level、filename。
- 不要输出 Markdown。
- 每个候选人仍然只输出 conclusion、strengths、gaps、suggestions。
- 优势必须来自已匹配技能，短板必须来自缺失技能或本地短板。
- 文风要像 API/LLM 工程岗位的筛选意见，简洁、具体、专业。
""".strip()


def call_batch_ai_summary(results: list[dict[str, Any]], jd_text: str) -> str:
    ai_config = load_ai_config()
    if not ai_config:
        return "未配置 AI 密钥、Base URL 或模型名，因此本次没有调用 AI。"

    user_prompt = build_batch_ai_prompt(results, jd_text)
    if ai_config["provider"] == "anthropic":
        raw_text = call_anthropic_summary(ai_config, user_prompt)
    else:
        raw_text = call_openai_compatible_summary(ai_config, user_prompt)

    try:
        summary, reviews = parse_ai_review_response(raw_text)
    except (json.JSONDecodeError, ValueError) as error:
        return f"智能复核返回格式不是结构化卡片，已保留本地规则结果：{error}"

    applied_count = apply_ai_reviews_to_results(results, reviews)
    if not applied_count:
        return "智能复核已返回结果，但没有匹配到候选人文件名；已保留本地规则卡片。"

    return f"{summary} 已更新 {applied_count} 份候选人卡片。"


def parse_ai_review_response(raw_text: str) -> tuple[str, dict[str, dict[str, Any]]]:
    text = raw_text.strip()
    if text.startswith("```"):
        text = re.sub(r"^```(?:json)?\s*", "", text)
        text = re.sub(r"\s*```$", "", text)

    start = text.find("{")
    end = text.rfind("}")
    if start >= 0 and end > start:
        text = text[start : end + 1]

    data = json.loads(text)
    if not isinstance(data, dict):
        raise ValueError("AI 返回不是 JSON 对象")

    summary = clean_ai_line(
        str(data.get("summary") or "AI 已按候选人卡片完成专业复核，分数和等级仍以本地规则为准。")
    )
    if len(summary) > 120:
        summary = f"{summary[:120]}..."
    reviews: dict[str, dict[str, Any]] = {}
    raw_reviews = data.get("reviews", [])
    if not isinstance(raw_reviews, list):
        return summary, reviews

    for raw_review in raw_reviews:
        if not isinstance(raw_review, dict):
            continue

        filename = str(raw_review.get("filename") or "").strip()
        if not filename:
            continue

        reviews[filename] = {
            "conclusion": normalize_ai_text(raw_review.get("conclusion")),
            "strengths": normalize_ai_lines(raw_review.get("strengths"), limit=4),
            "gaps": normalize_ai_lines(raw_review.get("gaps"), limit=4),
            "suggestions": normalize_ai_lines(raw_review.get("suggestions"), limit=4),
        }

    return summary, reviews


def normalize_ai_text(value: Any) -> str:
    if not isinstance(value, str):
        return ""
    return clean_ai_line(value)


def normalize_ai_lines(value: Any, limit: int = 4) -> list[str]:
    raw_lines: list[str] = []
    if isinstance(value, list):
        raw_lines = [str(item) for item in value]
    elif isinstance(value, str):
        raw_lines = re.split(r"[\r\n]+", value)

    lines = []
    for raw_line in raw_lines:
        line = clean_ai_line(raw_line)
        if line:
            lines.append(line)

    return lines[:limit]


def clean_ai_line(value: str) -> str:
    line = value.strip()
    line = re.sub(r"^[-*•\d.、\s]+", "", line)
    return line.strip()


def apply_ai_reviews_to_results(results: list[dict[str, Any]], reviews: dict[str, dict[str, Any]]) -> int:
    applied_count = 0
    for item in results:
        if not item.get("ok"):
            continue

        review = reviews.get(str(item.get("filename", "")))
        if not review:
            continue

        if review.get("conclusion"):
            item["conclusion"] = review["conclusion"]
        for key in ("strengths", "gaps", "suggestions"):
            lines = review.get(key)
            if lines:
                item[key] = lines

        item["ai_enhanced"] = True
        applied_count += 1

    return applied_count


def call_openai_compatible_summary(
    ai_config: dict[str, str],
    user_prompt: str,
    system_prompt: str = AI_SYSTEM_PROMPT,
) -> str:
    response = httpx.post(
        f"{ai_config['base_url']}/chat/completions",
        headers={
            "Authorization": f"Bearer {ai_config['api_key']}",
            "Content-Type": "application/json",
        },
        json={
            "model": ai_config["model"],
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "temperature": 0.3,
        },
        timeout=AI_TIMEOUT_SECONDS,
    )
    response.raise_for_status()
    data = response.json()
    return data["choices"][0]["message"]["content"] or ""


def call_anthropic_summary(
    ai_config: dict[str, str],
    user_prompt: str,
    system_prompt: str = AI_SYSTEM_PROMPT,
    max_tokens: int = AI_MAX_TOKENS,
) -> str:
    api_key = ai_config["api_key"]
    response = httpx.post(
        build_anthropic_messages_url(ai_config["base_url"]),
        headers={
            "x-api-key": api_key,
            "Authorization": f"Bearer {api_key}",
            "anthropic-version": "2023-06-01",
            "Content-Type": "application/json",
        },
        json={
            "model": ai_config["model"],
            "max_tokens": max_tokens,
            "temperature": 0.3,
            "system": system_prompt,
            "messages": [{"role": "user", "content": user_prompt}],
        },
        timeout=AI_TIMEOUT_SECONDS,
    )
    response.raise_for_status()
    data = response.json()
    parts = data.get("content", [])
    return "\n".join(part.get("text", "") for part in parts if part.get("type") == "text").strip()


def build_anthropic_messages_url(base_url: str) -> str:
    clean_url = base_url.rstrip("/")
    if clean_url.endswith("/v1/messages"):
        return clean_url
    if clean_url.endswith("/v1"):
        return f"{clean_url}/messages"
    return f"{clean_url}/v1/messages"


def build_ai_report_prompt(payload: dict[str, Any]) -> str:
    results = payload.get("results", [])
    compact_results = []
    for index, item in enumerate(results, start=1):
        if not item.get("ok"):
            compact_results.append(
                {
                    "rank": index,
                    "filename": item.get("filename"),
                    "ok": False,
                    "error": item.get("error"),
                }
            )
            continue

        compact_results.append(
            {
                "rank": index,
                "filename": item.get("filename"),
                "score": item.get("score"),
                "level": item.get("level"),
                "conclusion": item.get("conclusion"),
                "matched_skills": item.get("matched_skills"),
                "missing_skills": item.get("missing_skills"),
                "strengths": item.get("strengths", [])[:4],
                "gaps": item.get("gaps", [])[:4],
                "suggestions": item.get("suggestions", [])[:4],
            }
        )

    return f"""
分析批次：
- 生成时间：{payload.get("created_at", "未知")}
- JD 来源：{payload.get("jd_source", "未知")}
- 简历数量：{len(results)}

结构化匹配结果：
{json.dumps(compact_results, ensure_ascii=False, indent=2)}

请写一份“候选人批量筛选汇总报告”，结构如下：

# 候选人批量筛选汇总报告

## 1. 总体结论
用 2 到 4 句话概括这个候选人池是否适合当前 AI 应用开发岗位。

## 2. 推荐进入下一轮
列出建议进入下一轮的人，说明理由；如果只有一个人适合，要明确。

## 3. 候选人对比
按排名逐个说明每位候选人的定位、优势、主要风险和建议追问点。

## 4. 共性短板
总结候选人池在 API 接入、RAG、工具调用、向量数据库、Docker 或工程化方面的共性缺口。

## 5. 面试建议
给出 5 到 8 条技术面试追问，问题要具体，适合验证真实项目经验。

## 6. 招聘决策建议
给出可执行建议：优先推进、备选培养、暂缓或淘汰。

要求：
- 保留候选人原始分数和等级。
- 不要把读取失败的文件当成合格候选人。
- 不要虚构没有出现在结构化结果里的技能。
- 语言要专业、直接，像技术负责人写给招聘团队的报告。
""".strip()


def call_ai_report(payload: dict[str, Any]) -> str:
    ai_config = load_ai_config()
    if not ai_config:
        raise RuntimeError("未配置 AI 密钥、Base URL 或模型名，无法生成 AI 汇总报告。")

    prompt = build_ai_report_prompt(payload)
    if ai_config["provider"] == "anthropic":
        return call_anthropic_summary(
            ai_config,
            prompt,
            system_prompt=AI_REPORT_SYSTEM_PROMPT,
            max_tokens=AI_REPORT_MAX_TOKENS,
        )
    return call_openai_compatible_summary(ai_config, prompt, system_prompt=AI_REPORT_SYSTEM_PROMPT)


def write_ai_report_docx(report_text: str, payload: dict[str, Any], path: Path) -> None:
    document = Document()
    document.add_heading("AI 候选人汇总报告", level=1)
    document.add_paragraph(f"生成时间：{datetime.now().isoformat(timespec='seconds')}")
    document.add_paragraph(f"原始分析时间：{payload.get('created_at', '未知')}")
    document.add_paragraph(f"JD 来源：{payload.get('jd_source', '未知')}")
    document.add_paragraph(f"简历数量：{len(payload.get('results', []))}")

    for raw_line in report_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith(("- ", "* ")):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)

    document.save(path)


def load_saved_payload(json_filename: str) -> dict[str, Any]:
    safe_name = safe_filename(json_filename)
    if Path(safe_name).suffix.lower() != ".json":
        raise HTTPException(status_code=400, detail="只能基于 JSON 结果文件生成 AI 报告")

    path = RESULTS_DIR / safe_name
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=404, detail="结果 JSON 文件不存在")

    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as error:
        raise HTTPException(status_code=400, detail=f"结果 JSON 无法解析：{error}") from error

    if not isinstance(payload, dict) or not isinstance(payload.get("results"), list):
        raise HTTPException(status_code=400, detail="结果 JSON 格式不正确")

    return payload


def get_level(match_data: dict[str, Any]) -> str:
    diagnosis = match_data.get("diagnosis", {})
    return str(diagnosis.get("level", "未知"))


def serialize_agent_trace(trace: list[AgentTraceStep]) -> list[dict[str, str]]:
    return [
        {
            "phase": step.phase,
            "thought": step.thought,
            "tool": step.tool,
            "action": step.action,
            "observation": step.observation,
        }
        for step in trace
    ]


def build_web_rag_query(filename: str, match_data: dict[str, Any]) -> str:
    return build_rag_query(
        f"分析 {filename} 与岗位 JD 的匹配证据",
        match_data,
    )


def build_result_item(
    filename: str,
    resume_text: str,
    jd_text: str,
    memory_tool: JobMemoryTool | None = None,
    rag_tool: JobRAGTool | None = None,
    batch_index: int = 1,
) -> dict[str, Any]:
    trace: list[AgentTraceStep] = [
        build_agent_trace_step(
            "Perception",
            "接收 Web 上传的简历和岗位 JD，并转换成可分析文本。",
            "load_text_file",
            "读取上传文件",
            f"简历 {filename}，文本 {len(resume_text)} 字；JD 文本 {len(jd_text)} 字。",
        )
    ]

    memory_context = ""
    if memory_tool:
        memory_summary = memory_tool.run({"action": "summary", "limit": 3})
        memory_matches = memory_tool.run({"action": "search", "query": f"{filename} {jd_text[:300]}", "limit": 3})
        memory_context = f"记忆摘要：{memory_summary}\n相关历史：{memory_matches}"
        trace.append(
            build_agent_trace_step(
                "Memory",
                "读取用户历史求职画像和过往分析短板，为本轮判断提供连续性。",
                "JobMemoryTool",
                "summary + search",
                truncate_text(memory_context),
            )
        )

    if rag_tool:
        jd_observation = rag_tool.run(
            {
                "action": "add_document",
                "source": f"current_jd_for_{filename}",
                "doc_type": "jd",
                "text": jd_text,
                "metadata": {"filename": filename},
            }
        )
        resume_observation = rag_tool.run(
            {
                "action": "add_document",
                "source": filename,
                "doc_type": "resume",
                "text": resume_text,
                "metadata": {"batch_index": batch_index},
            }
        )
        trace.append(
            build_agent_trace_step(
                "RAG Indexing",
                "把当前简历和 JD 分块写入轻量检索索引，用于后续证据召回。",
                "JobRAGTool",
                "add_document(jd) + add_document(resume)",
                f"{jd_observation} {resume_observation}",
            )
        )

    candidate_skills = extract_candidate_skills(resume_text)
    job_requirements = extract_job_requirements(jd_text)
    match_result = compute_match_score(candidate_skills, job_requirements)
    match_data = json.loads(match_result)
    score = int(match_data.get("match_score", 0))

    trace.extend(
        [
            build_agent_trace_step(
                "Thought",
                "抽取候选人技能和岗位要求，形成结构化对比对象。",
                "extract_candidate_skills / extract_job_requirements",
                "抽取技能画像与岗位要求",
                truncate_text(f"候选人技能：{candidate_skills}\n岗位要求：{job_requirements}"),
            ),
            build_agent_trace_step(
                "Planning + Tool Selection",
                "选择匹配评分工具，根据权重、类别覆盖和缺口项生成可解释分数。",
                "compute_match_score",
                "计算匹配分、优势、短板和建议",
                truncate_text(match_result),
            ),
        ]
    )

    rag_context = ""
    if rag_tool:
        rag_context = rag_tool.run({"action": "search", "query": build_web_rag_query(filename, match_data), "limit": 4})
        trace.append(
            build_agent_trace_step(
                "RAG Retrieval",
                "检索简历和 JD 中最能支撑当前匹配结论的证据片段。",
                "JobRAGTool",
                "search(query)",
                truncate_text(rag_context),
            )
        )

    if memory_tool:
        memory_write = memory_tool.run(
            {
                "action": "add",
                "memory_type": "episodic",
                "content": build_gap_memory_content(filename, match_data),
                "importance": 0.8,
                "metadata": {
                    "resume_label": filename,
                    "score": score,
                    "level": get_level(match_data),
                    "missing_skills": join_skill_names(match_data.get("missing_skills", [])),
                    "matched_skills": join_skill_names(match_data.get("matched_skills", [])),
                    "target_role": "AI 应用开发 / Python 后端",
                },
            }
        )
        trace.append(
            build_agent_trace_step(
                "Memory Write",
                "把本次候选人分析写入情景记忆，便于后续追踪简历版本和能力短板。",
                "JobMemoryTool",
                "add(episodic)",
                memory_write,
            )
        )

    result = {
        "filename": filename,
        "ok": True,
        "score": score,
        "level": get_level(match_data),
        "conclusion": get_score_comment(score),
        "strengths": build_strength_lines(match_data),
        "gaps": build_gap_lines(match_data),
        "suggestions": build_suggestion_lines(match_data),
        "matched_skills": join_skill_names(match_data.get("matched_skills", [])),
        "missing_skills": join_skill_names(match_data.get("missing_skills", [])),
        "agent_trace": serialize_agent_trace(trace),
        "memory_context": memory_context,
        "rag_context": rag_context,
        "raw": match_data,
    }

    return result


def append_markdown_agent_details(lines: list[str], item: dict[str, Any]) -> None:
    trace = item.get("agent_trace") or []
    if trace:
        lines.extend(["", "### Agent 循环", ""])
        for step in trace:
            lines.extend(
                [
                    f"- **{step.get('phase', 'Unknown')}**",
                    f"  - Thought：{step.get('thought', '')}",
                    f"  - Tool Selection：{step.get('tool', '')}",
                    f"  - Action：{step.get('action', '')}",
                    f"  - Observation：{step.get('observation', '')}",
                ]
            )

    if item.get("memory_context"):
        lines.extend(["", "### Memory 记忆上下文", "", item["memory_context"], ""])

    if item.get("rag_context"):
        lines.extend(["", "### RAG 检索上下文", "", item["rag_context"], ""])


def append_docx_agent_details(document: Document, item: dict[str, Any]) -> None:
    trace = item.get("agent_trace") or []
    if trace:
        document.add_heading("Agent 循环", level=3)
        for step in trace:
            document.add_paragraph(str(step.get("phase", "Unknown")), style="List Bullet")
            document.add_paragraph(f"Thought：{step.get('thought', '')}")
            document.add_paragraph(f"Tool Selection：{step.get('tool', '')}")
            document.add_paragraph(f"Action：{step.get('action', '')}")
            document.add_paragraph(f"Observation：{step.get('observation', '')}")

    if item.get("memory_context"):
        document.add_heading("Memory 记忆上下文", level=3)
        document.add_paragraph(str(item["memory_context"]))

    if item.get("rag_context"):
        document.add_heading("RAG 检索上下文", level=3)
        document.add_paragraph(str(item["rag_context"]))


def write_markdown_report(payload: dict[str, Any], path: Path) -> None:
    lines = [
        "# 简历 JD 匹配分析结果",
        "",
        f"- 生成时间：{payload['created_at']}",
        f"- JD 来源：{payload['jd_source']}",
        f"- 简历数量：{len(payload['results'])}",
        "",
    ]

    for index, item in enumerate(payload["results"], start=1):
        lines.extend([f"## {index}. {item['filename']}", ""])
        if not item.get("ok"):
            lines.extend([f"读取失败：{item.get('error', '未知错误')}", ""])
            continue

        lines.extend(
            [
                f"- 匹配分：{item['score']} / 100",
                f"- 匹配等级：{item['level']}",
                f"- 一句话结论：{item['conclusion']}",
                f"- 已匹配技能：{item['matched_skills']}",
                f"- 缺失技能：{item['missing_skills']}",
                "",
                "### 优势",
                "",
            ]
        )
        lines.extend(f"- {line}" for line in item["strengths"])
        lines.extend(["", "### 短板", ""])
        lines.extend(f"- {line}" for line in item["gaps"])
        lines.extend(["", "### 改进建议", ""])
        lines.extend(f"- {line}" for line in item["suggestions"])
        append_markdown_agent_details(lines, item)
        lines.append("")

    if payload.get("ai_summary"):
        lines.extend(["## 智能结构化复核", "", payload["ai_summary"], ""])

    path.write_text("\n".join(lines), encoding="utf-8")


def write_docx_report(payload: dict[str, Any], path: Path) -> None:
    document = Document()
    document.add_heading("简历 JD 匹配分析结果", level=1)
    document.add_paragraph(f"生成时间：{payload['created_at']}")
    document.add_paragraph(f"JD 来源：{payload['jd_source']}")
    document.add_paragraph(f"简历数量：{len(payload['results'])}")
    document.add_paragraph(f"AI 补充分析：{'已启用' if payload['ai_enabled'] else '未启用'}")
    if payload.get("ai_summary"):
        document.add_heading("智能结构化复核", level=2)
        document.add_paragraph(payload["ai_summary"])

    for index, item in enumerate(payload["results"], start=1):
        document.add_heading(f"{index}. {item['filename']}", level=2)
        if not item.get("ok"):
            document.add_paragraph(f"读取失败：{item.get('error', '未知错误')}")
            continue

        document.add_paragraph(f"匹配分：{item['score']} / 100")
        document.add_paragraph(f"匹配等级：{item['level']}")
        document.add_paragraph(f"一句话结论：{item['conclusion']}")
        document.add_paragraph(f"已匹配技能：{item['matched_skills']}")
        document.add_paragraph(f"缺失技能：{item['missing_skills']}")

        document.add_heading("优势", level=3)
        for line in item["strengths"]:
            document.add_paragraph(line, style="List Bullet")

        document.add_heading("短板", level=3)
        for line in item["gaps"]:
            document.add_paragraph(line, style="List Bullet")

        document.add_heading("改进建议", level=3)
        for line in item["suggestions"]:
            document.add_paragraph(line, style="List Bullet")

        if item.get("ai_summary"):
            document.add_heading("AI 补充分析", level=3)
            document.add_paragraph(item["ai_summary"])

        append_docx_agent_details(document, item)

    document.save(path)


def write_excel_ranking(results: list[dict[str, Any]], path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "排名"
    headers = ["排名", "人物/文件名", "匹配分", "匹配等级"]
    sheet.append(headers)

    header_fill = PatternFill("solid", fgColor="1F7A5A")
    header_font = Font(color="FFFFFF", bold=True)
    for cell in sheet[1]:
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    rank = 1
    for item in results:
        if not item.get("ok"):
            continue
        sheet.append([rank, item["filename"], item["score"], item["level"]])
        rank += 1

    widths = [10, 28, 12, 16]
    for column_index, width in enumerate(widths, start=1):
        sheet.column_dimensions[chr(64 + column_index)].width = width

    for row in sheet.iter_rows(min_row=2):
        row[0].alignment = Alignment(horizontal="center")
        row[2].alignment = Alignment(horizontal="center")
        row[3].alignment = Alignment(horizontal="center")

    workbook.save(path)


def save_analysis_files(
    results: list[dict[str, Any]],
    jd_source: str,
    ai_enabled: bool,
    ai_summary: str | None,
) -> dict[str, str]:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    payload = {
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "jd_source": jd_source,
        "ai_enabled": ai_enabled,
        "ai_summary": ai_summary,
        "results": results,
    }
    json_path = RESULTS_DIR / f"match_results_{timestamp}.json"
    md_path = RESULTS_DIR / f"match_results_{timestamp}.md"
    docx_path = RESULTS_DIR / f"match_results_{timestamp}.docx"
    xlsx_path = RESULTS_DIR / f"match_ranking_{timestamp}.xlsx"

    json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    write_markdown_report(payload, md_path)
    write_docx_report(payload, docx_path)
    write_excel_ranking(results, xlsx_path)

    return {
        "json": json_path.name,
        "markdown": md_path.name,
        "docx": docx_path.name,
        "excel": xlsx_path.name,
    }


def sort_results(results: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return sorted(results, key=lambda item: item.get("score", -1), reverse=True)


@app.get("/health")
async def health() -> dict[str, Any]:
    return {
        "app": APP_SLUG,
        "name": APP_NAME,
        "project_root": str(ROOT_DIR),
        "memory_rag_enabled": True,
    }


@app.get("/", response_class=HTMLResponse)
async def index(request: Request) -> HTMLResponse:
    return templates.TemplateResponse(
        "pages/index.html",
        {
            "request": request,
            "results": None,
            "saved_files": None,
            "message": "",
            "ai_enabled": False,
        },
    )


@app.post("/analyze", response_class=HTMLResponse)
async def analyze(
    request: Request,
    jd_text: str = Form(""),
    enable_ai: bool = Form(False),
    jd_file: UploadFile | None = File(None),
    resume_files: list[UploadFile] = File(...),
) -> HTMLResponse:
    messages: list[str] = []
    results: list[dict[str, Any]] = []
    resume_texts: dict[str, str] = {}
    user_id = get_optional_env("JOB_FIT_USER_ID", "WEB_USER_ID") or DEFAULT_WEB_USER_ID
    memory_tool = JobMemoryTool(user_id=user_id)
    rag_tool = JobRAGTool(namespace=f"web_{user_id}")
    rag_tool.run({"action": "clear"})
    messages.append("已启用轻量 Memory 与 RAG：本轮结果会写入求职画像记忆，并从当前 JD / 简历检索证据。")

    with tempfile.TemporaryDirectory() as temp_dir_name:
        temp_dir = Path(temp_dir_name)

        jd_content = jd_text.strip()
        jd_source = "页面粘贴 JD"
        if jd_file and jd_file.filename:
            jd_path = save_upload(jd_file, temp_dir)
            jd_content = load_text_file(jd_path)
            jd_source = safe_filename(jd_file.filename)
        elif not jd_content:
            jd_content = get_jd_text(None)
            jd_source = "内置示例 JD"
            messages.append("未填写 JD，已使用内置示例 JD。")

        for upload in resume_files:
            if not upload.filename:
                continue

            filename = safe_filename(upload.filename)
            try:
                resume_path = save_upload(upload, temp_dir)
                resume_text = load_text_file(resume_path)
                resume_texts[filename] = resume_text
                results.append(
                    build_result_item(
                        filename,
                        resume_text,
                        jd_content,
                        memory_tool=memory_tool,
                        rag_tool=rag_tool,
                        batch_index=len(results) + 1,
                    )
                )
            except Exception as error:
                results.append(
                    {
                        "filename": filename,
                        "ok": False,
                        "error": str(error),
                    }
                )

    ranked_results = sort_results(results)
    ai_summary = None
    if enable_ai:
        try:
            ai_summary = call_batch_ai_summary(ranked_results, jd_content)
        except httpx.TimeoutException:
            ai_summary = f"AI 批量分析超时：等待超过 {AI_TIMEOUT_SECONDS} 秒，已保留本地规则结果。"
        except Exception as error:
            ai_summary = f"AI 批量分析失败：{error}"

    saved_files = save_analysis_files(ranked_results, jd_source, enable_ai, ai_summary)
    try:
        saved_count = save_analysis_results(
            jd_source=jd_source,
            jd_text=jd_content,
            results=ranked_results,
            resume_texts=resume_texts,
        )
        messages.append(f"已写入 SQLite 数据库 {saved_count} 条成功匹配记录。")
    except Exception as error:
        messages.append(f"SQLite 持久化失败，已保留导出文件：{error}")

    return templates.TemplateResponse(
        "pages/index.html",
        {
            "request": request,
            "results": ranked_results,
            "saved_files": saved_files,
            "message": " ".join(messages),
            "ai_enabled": enable_ai,
            "ai_summary": ai_summary,
        },
    )


@app.get("/results/{filename}")
async def download_result(filename: str) -> FileResponse:
    safe_name = safe_filename(filename)
    path = RESULTS_DIR / safe_name
    if not path.exists() or not path.is_file():
        raise HTTPException(status_code=404, detail="结果文件不存在")

    media_types = {
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".json": "application/json",
        ".md": "text/markdown; charset=utf-8",
    }
    return FileResponse(path, filename=safe_name, media_type=media_types.get(path.suffix.lower()))


@app.get("/ai-report/{json_filename}")
async def generate_ai_report(json_filename: str) -> JSONResponse:
    safe_name = safe_filename(json_filename)
    payload = load_saved_payload(safe_name)
    source_stem = Path(safe_name).stem
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_stem = f"ai_summary_report_{source_stem}_{timestamp}"
    md_path = RESULTS_DIR / f"{report_stem}.md"

    try:
        report_text = call_ai_report(payload)
    except httpx.TimeoutException as error:
        raise HTTPException(status_code=504, detail=f"AI 汇总报告生成超时：{error}") from error
    except Exception as error:
        raise HTTPException(status_code=500, detail=f"AI 汇总报告生成失败：{error}") from error

    md_path.write_text(report_text, encoding="utf-8")
    return JSONResponse({"report": report_text, "markdown": md_path.name})


def port_is_free(port: int, host: str = DEFAULT_HOST) -> bool:
    import socket

    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        return sock.connect_ex((host, port)) != 0


def find_available_port(start_port: int = DEFAULT_PORT, host: str = DEFAULT_HOST) -> int:
    port = start_port
    while not port_is_free(port, host):
        port += 1
    return port


if __name__ == "__main__":
    import uvicorn

    requested_port = int(os.getenv("WEB_APP_PORT", str(DEFAULT_PORT)))
    port = find_available_port(requested_port)
    url = f"http://{DEFAULT_HOST}:{port}"
    if port != requested_port:
        print(f"{APP_NAME} requested port {requested_port} is busy; using {url}", flush=True)
    else:
        print(f"{APP_NAME} is running at {url}", flush=True)
    print(f"Health check: {url}/health", flush=True)
    uvicorn.run(app, host=DEFAULT_HOST, port=port)
